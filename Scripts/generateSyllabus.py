from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_UNDERLINE
import openpyxl, os
import pandas as pd

document = Document()
section = document.sections[0]
header = section.header

f = open( "q&a.txt", "r")
g = f.read()
answers = [s.replace('"', '').replace("[", "").replace("]", "") for s in g.split(',')]

# Load schedule excel
df = pd.read_excel('C:/Users/mahya/Desktop/Capstone/courseorganizer-main/Scripts/test.xlsx', sheet_name='Sheet1')

# Template of the syllabus
paragraph = header.paragraphs[0]
paragraph.text = "Course Syllabus\t\t" + answers[7]
paragraph.style = document.styles["Header"]

title = document.add_paragraph(answers[0])
title_format = title.paragraph_format
title_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

title2 = document.add_paragraph(answers[1])
title2_format = title2.paragraph_format
title2_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

title3 = document.add_paragraph(answers[3])
title3_format = title3.paragraph_format
title3_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

records = [
    ('Class Location: ' + answers[4], 'E-mail: ' + answers[50]),
    ('Office location: ' + answers[7], 'Office hours: ' + answers[6]),
    
]

if (answers[14].upper() == 'Yes'):
    records.append('Phone: ' + answers[14])
if (answers[20].upper() == 'Yes'):
    records.append('Lab location: ' + answers[51])
if (answers[22].upper() == 'Yes'):
    records.append(('Discussion seminar time: ' + answers[52], 'Discussion Seminar location: ' + answers[53]))
if (answers[16].upper() == 'Yes'):
    records.append('Virtual office hours(Zoom)\nMeeting ID: '+ answers[17] + '\nPasscode: \n' + answers[18])

table = document.add_table(rows=1, cols=2, style = "Table Grid")
row = table.rows[0].cells
row[0].text = 'Instructor: ' + answers[2]
row[1].text = 'Class Time: ' + answers[5]
for x, y in records:
    row_cells = table.add_row().cells
    row_cells[0].text = x
    row_cells[1].text = y

paragraph2 = document.add_paragraph('\n')

if (answers[28].upper() == 'YES'):
    run = paragraph2.add_run('Prerequisites')
    run.underline = True
    run.underline = WD_UNDERLINE.SINGLE
    paragraph2.add_run('\n' + answers[29])

paragraph2.add_run('\n')
paragraph2.add_run('\nCourse Description').underline = WD_UNDERLINE.SINGLE
paragraph2.add_run('\n' + answers[12])

if (answers[30].upper() == 'YES'):
    paragraph2.add_run('\n')
    paragraph2.add_run('\nCourse Materials').underline = WD_UNDERLINE.SINGLE
    paragraph2.add_run('\n' 'Textbook: ' + answers[31])
    paragraph2.add_run('\n' 'ISBN: ' + answers[32])

if (answers[33].upper() == 'YES'):
    paragraph2.add_run('\n')
    paragraph2.add_run('\nLearning Objectives').underline = WD_UNDERLINE.SINGLE
    paragraph2.add_run('\nStudents who successfully complete this course will achieve the following learning objectives:')
    paragraph2.add_run('\n' + answers[34])

if (answers[20].upper() == 'YES'):
    paragraph2.add_run('\n')
    paragraph2.add_run('\nLab Policy').underline = WD_UNDERLINE.SINGLE
    paragraph2.add_run('\n' + answers[21])

if (answers[35].upper() == 'YES'):
    paragraph2.add_run('\n')
    paragraph2.add_run('\nAssignments').underline = WD_UNDERLINE.SINGLE
    paragraph2.add_run('\n' + answers[36])

paragraph2.add_run('\n')
paragraph2.add_run('\nExpectations').underline = WD_UNDERLINE.SINGLE
paragraph2.add_run('\n' + answers[37])

if (answers[38].upper() == 'YES'):
    paragraph2.add_run('\n')
    paragraph2.add_run('\nQuiz').underline = WD_UNDERLINE.SINGLE
    paragraph2.add_run('\n' + answers[39])

if (answers[40].upper() == 'YES'):
    paragraph2.add_run('\n')
    paragraph2.add_run('\nExam').underline = WD_UNDERLINE.SINGLE
    paragraph2.add_run('\n' + answers[41])

if (answers[22].upper() == 'YES'):
    paragraph2.add_run('\n')
    paragraph2.add_run('\nDiscussion Policy').underline = WD_UNDERLINE.SINGLE
    paragraph2.add_run('\n' + answers[23])


paragraph2.add_run('\n')
paragraph2.add_run('\nAttendance').underline = WD_UNDERLINE.SINGLE
paragraph2.add_run('\n' + answers[11])

paragraph2.add_run('\n')
paragraph2.add_run('\nGrading').underline = WD_UNDERLINE.SINGLE
paragraph2.add_run('\n' + answers[43])

paragraph2.add_run('\n')
paragraph2.add_run('\nDisability Services').underline = WD_UNDERLINE.SINGLE
paragraph2.add_run('\n' + answers[10])

paragraph2.add_run('\n')
paragraph2.add_run('\nHonor Code').underline = WD_UNDERLINE.SINGLE
paragraph2.add_run('\n' + answers[9])

if (answers[44].upper() == 'YES'):
    paragraph2.add_run('\n')
    paragraph2.add_run('\nOnline Resources').underline = WD_UNDERLINE.SINGLE
    paragraph2.add_run('\n' + answers[45])

if (answers[46].upper() == 'YES'):
    paragraph2.add_run('\n')
    paragraph2.add_run('\nExtra Credit').underline = WD_UNDERLINE.SINGLE
    paragraph2.add_run('\n' + answers[47])

paragraph2.add_run('\n')
paragraph2.add_run('\nFinal Exam').underline = WD_UNDERLINE.SINGLE
paragraph2.add_run('\n' + answers[48])

paragraph2.add_run('\n')
paragraph2.add_run('\nInclement Weather').underline = WD_UNDERLINE.SINGLE
paragraph2.add_run('\n' + answers[13])

paragraph2.add_run('\n')
paragraph2.add_run('\nWithdrawals').underline = WD_UNDERLINE.SINGLE
paragraph2.add_run('\n' + answers[49])

if (answers[27].upper() == 'YES'):
    paragraph2.add_run('\n')
    paragraph2.add_run('\nSyllabus Update').underline = WD_UNDERLINE.SINGLE
    paragraph2.add_run('\nThis syllabus can be updated at any point in time in the semester')

paragraph2.add_run('\n')
paragraph2.add_run('\nLecture Schedule').underline = WD_UNDERLINE.SINGLE

#schedule table
df.fillna('', inplace=True)
table2 = document.add_table(rows=1, cols=len(df.columns), style = "Table Grid")

# Add header row to the table
hdr_cells = table2.rows[0].cells
for i in range(len(df.columns)):
    hdr_cells[i].text = df.columns[i]

# Add data to the table
for index, row in df.iterrows():
    row_cells = table2.add_row().cells
    for i in range(len(df.columns)):
        row_cells[i].text = str(row[i])

document.save('syllabus.docx')
